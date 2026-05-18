#!/usr/bin/env python3
import sys
import re
from docx import Document

def extract_settings_json_config(doc_path):
    """提取settings.json配置内容"""
    try:
        doc = Document(doc_path)
        settings_content = []
        in_settings_section = False
        settings_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # 搜索settings.json相关内容
            if re.search(r'settings\.json|添加以下两行', text):
                in_settings_section = True
                settings_content.append(f"找到配置内容: {text}")
                
            elif in_settings_section:
                # 收集配置行
                if text and re.search(r'\{|"ai_assistant|"ppe_data_label_trae', text):
                    settings_lines.append(text)
                elif text and re.search(r'\}|reload|window', text, re.IGNORECASE):
                    # 可能是配置结束
                    settings_content.append("配置内容:")
                    settings_content.extend(settings_lines)
                    break
        
        # 如果没有找到具体配置，尝试提取更多上下文
        if not settings_lines:
            settings_content.append("\n搜索更多配置信息...")
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if re.search(r'json|配置|ppe|ai_assistant', text, re.IGNORECASE):
                    settings_content.append(f"[{i+1}] {text}")
        
        return "\n".join(settings_content)
        
    except Exception as e:
        return f"提取配置时出错: {str(e)}"

def extract_appendix_e_details(doc_path):
    """提取附录E的详细步骤"""
    try:
        doc = Document(doc_path)
        appendix_content = []
        in_appendix_e = False
        step_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # 查找附录E
            if re.search(r'附录\s*[Ee]:.*Trea.*PPE', text):
                in_appendix_e = True
                appendix_content.append(f"# {text}")
                continue
                
            # 如果已经在附录E中
            if in_appendix_e:
                # 检查是否到了下一个附录
                if re.search(r'附录\s*[Ff]|^[0-9]+\.[0-9]+', text) and step_count > 0:
                    break
                    
                # 提取步骤
                if re.search(r'^[0-9]+\.', text):
                    step_count += 1
                    appendix_content.append(f"\n步骤{step_count}: {text}")
                elif text:
                    appendix_content.append(f"  {text}")
        
        return "\n".join(appendix_content)
        
    except Exception as e:
        return f"提取附录E详情时出错: {str(e)}"

def main():
    doc_path = "/app/logsift/docs/环境 & prompt & traj 标注标准.docx"
    
    print("="*80)
    print("提取settings.json配置内容")
    print("="*80)
    
    settings_config = extract_settings_json_config(doc_path)
    print(settings_config)
    
    print("\n" + "="*80)
    print("提取附录E详细步骤")
    print("="*80)
    
    appendix_details = extract_appendix_e_details(doc_path)
    print(appendix_details)
    
    # 保存重要信息
    with open("/app/logsift/.trae/ppe_settings_info.txt", "w", encoding="utf-8") as f:
        f.write("PPE配置信息汇总\n")
        f.write("="*60 + "\n\n")
        f.write(settings_config + "\n\n")
        f.write(appendix_details)

if __name__ == "__main__":
    main()