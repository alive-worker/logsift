#!/usr/bin/env python3
import sys
import re
from docx import Document

def extract_full_document(doc_path):
    """提取完整文档内容"""
    try:
        doc = Document(doc_path)
        full_content = []
        
        print(f"文档段落数: {len(doc.paragraphs)}")
        
        # 提取所有段落
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # 检查样式信息
                style = para.style.name if para.style else "无样式"
                full_content.append(f"[段落 {i+1}, 样式: {style}] {text}")
        
        return "\n".join(full_content)
        
    except Exception as e:
        return f"提取文档时出错: {str(e)}"

def search_ppe_config(doc_path):
    """搜索PPE配置相关内容"""
    try:
        doc = Document(doc_path)
        ppe_content = []
        in_ppe_section = False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # 搜索PPE相关关键词
            if re.search(r'PPE|ppe|Prompt.*Processing.*Environment', text, re.IGNORECASE):
                ppe_content.append(f"找到PPE相关内容: {text}")
                in_ppe_section = True
            elif in_ppe_section:
                # 继续收集直到遇到新的标题
                if re.search(r'^[0-9]+\.|^附录|^第[一二三四五六七八九十]+章', text):
                    in_ppe_section = False
                elif text:
                    ppe_content.append(text)
        
        return "\n".join(ppe_content)
        
    except Exception as e:
        return f"搜索PPE配置时出错: {str(e)}"

def main():
    doc_path = "/app/logsift/docs/环境 & prompt & traj 标注标准.docx"
    
    print("="*80)
    print("搜索PPE配置相关内容...")
    print("="*80)
    
    ppe_config = search_ppe_config(doc_path)
    print(ppe_config)
    
    print("\n" + "="*80)
    print("提取文档前50个段落...")
    print("="*80)
    
    # 提取前50个段落
    try:
        doc = Document(doc_path)
        for i in range(min(50, len(doc.paragraphs))):
            text = doc.paragraphs[i].text.strip()
            if text:
                print(f"[{i+1}] {text}")
    except Exception as e:
        print(f"提取段落时出错: {str(e)}")

if __name__ == "__main__":
    main()